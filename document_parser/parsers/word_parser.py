from __future__ import annotations

from pathlib import Path
from typing import Any

from ..base_parser import BaseParser
from ..models import ParseResult


class WordParser(BaseParser):
    """Word 文档解析器（.docx 格式）。"""

    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: Path) -> ParseResult:
        """解析 Word 文档。

        Args:
            file_path: 要解析的文件路径

        Returns:
            ParseResult: 包含内容、元信息和摘要的解析结果
        """
        try:
            from docx import Document
        except ImportError:
            return ParseResult.from_error(
                "缺少依赖: python-docx。请运行 'pip install python-docx' 安装。",
                file_path=file_path,
            )

        try:
            document = Document(str(file_path))
        except Exception as e:
            return ParseResult.from_error(f"无法打开 Word 文档: {e}", file_path=file_path)

        content_parts = []
        metadata = self._extract_metadata(file_path, document)

        title = self._extract_title(document)
        if title:
            content_parts.append(f"# {title}\n")

        for element in document.element.body:
            if element.tag.endswith("p"):
                para = self._find_paragraph_by_element(document, element)
                if para:
                    text = self._process_paragraph(para)
                    if text:
                        content_parts.append(text)
            elif element.tag.endswith("tbl"):
                table = self._find_table_by_element(document, element)
                if table:
                    table_md = self._process_table(table)
                    if table_md:
                        content_parts.append(table_md)

        content = "\n".join(content_parts)

        summary = self._generate_summary(document, content)

        return ParseResult(
            content=content,
            metadata=metadata,
            summary=summary,
        )

    def _find_paragraph_by_element(self, document, element):
        """根据 XML 元素查找段落对象。"""
        from docx.text.paragraph import Paragraph

        for para in document.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_by_element(self, document, element):
        """根据 XML 元素查找表格对象。"""
        from docx.table import Table

        for table in document.tables:
            if table._element is element:
                return table
        return None

    def _extract_title(self, document) -> str | None:
        """提取文档标题（通常为第一个标题样式的段落）。

        Args:
            document: python-docx Document 对象

        Returns:
            str | None: 文档标题
        """
        for para in document.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                return para.text.strip()
            if style_name == "Title":
                return para.text.strip()
        return None

    def _process_paragraph(self, para) -> str:
        """处理段落，保留格式信息。

        Args:
            para: 段落对象

        Returns:
            str: 处理后的文本
        """
        text = para.text.strip()
        if not text:
            return ""

        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
                level = min(max(level, 1), 6)
                return f"\n{'#' * level} {text}\n"
            except ValueError:
                pass

        if style_name == "Title":
            return f"# {text}\n"

        if style_name == "List Bullet" or style_name.startswith("List"):
            return f"- {text}"

        return text

    def _process_table(self, table) -> str:
        """将表格转换为 Markdown 格式。

        Args:
            table: 表格对象

        Returns:
            str: Markdown 格式的表格
        """
        if not table.rows:
            return ""

        rows_data = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                cells.append(cell_text)
            rows_data.append(cells)

        if not rows_data:
            return ""

        max_cols = max(len(row) for row in rows_data)
        normalized_rows = []
        for row in rows_data:
            while len(row) < max_cols:
                row.append("")
            normalized_rows.append(row)

        header = normalized_rows[0]
        body = normalized_rows[1:] if len(normalized_rows) > 1 else []

        md_lines = []
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        for row in body:
            md_lines.append("| " + " | ".join(row) + " |")

        return "\n".join(md_lines) + "\n"

    def _extract_metadata(self, file_path: Path, document) -> dict[str, Any]:
        """提取 Word 文档的元信息。

        Args:
            file_path: 文件路径
            document: python-docx Document 对象

        Returns:
            dict: 元信息字典
        """
        stat = file_path.stat()

        core_props = document.core_properties

        paragraph_count = len(document.paragraphs)
        table_count = len(document.tables)

        word_count = 0
        char_count = 0
        for para in document.paragraphs:
            text = para.text
            word_count += len(text.split())
            char_count += len(text)

        heading_count = sum(
            1 for para in document.paragraphs
            if para.style and para.style.name and para.style.name.startswith("Heading")
        )

        metadata = {
            "file_name": file_path.name,
            "file_size": stat.st_size,
            "created_time": stat.st_ctime,
            "modified_time": stat.st_mtime,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "word_count": word_count,
            "char_count": char_count,
            "heading_count": heading_count,
        }

        if core_props.title:
            metadata["doc_title"] = core_props.title
        if core_props.author:
            metadata["doc_author"] = core_props.author
        if core_props.subject:
            metadata["doc_subject"] = core_props.subject
        if core_props.keywords:
            metadata["doc_keywords"] = core_props.keywords
        if core_props.created:
            metadata["doc_created"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["doc_modified"] = core_props.modified.isoformat()
        if core_props.last_modified_by:
            metadata["doc_last_modified_by"] = core_props.last_modified_by

        return metadata

    def _generate_summary(self, document, content: str) -> str:
        """生成 Word 文档摘要。

        Args:
            document: python-docx Document 对象
            content: 解析后的内容

        Returns:
            str: 摘要文本
        """
        title = self._extract_title(document)

        paragraph_count = len(document.paragraphs)
        table_count = len(document.tables)

        heading_count = sum(
            1 for para in document.paragraphs
            if para.style and para.style.name and para.style.name.startswith("Heading")
        )

        if title:
            return f"Word 文档，标题: {title}，共 {paragraph_count} 个段落，{heading_count} 个标题，{table_count} 个表格"
        else:
            first_para = ""
            for para in document.paragraphs:
                if para.text.strip():
                    first_para = para.text.strip()[:100]
                    break

            if first_para:
                return f"Word 文档，无标题，首段: {first_para}...，共 {paragraph_count} 个段落，{table_count} 个表格"
            return f"Word 文档，共 {paragraph_count} 个段落，{table_count} 个表格"